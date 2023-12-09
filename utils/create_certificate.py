# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.shared import RGBColor
# from docx2pdf import convert
#
# import os
#
#
# def create_certificate(user_id, test_id, fullname, science, class_num, date):
#     # Load the Word document
#     doc = Document("data/images/certificate.docx")
#
#     # Define the text to replace and the replacement text
#     text_to_find1 = "fullname"
#     text_to_find2 = "science fani class_num-sinf testini muvaffaqiyatli"
#     text_to_find3 = "date_time"
#
#     text2 = text_to_find2.replace('science', science).replace('class_num', f"{class_num}")
#
#     # Iterate through paragraphs in the document
#     for paragraph in doc.paragraphs:
#         if text_to_find1 in paragraph.text:
#             # Replace the text with the new value
#             paragraph.text = paragraph.text.replace(text_to_find1, fullname)
#
#             # Format the text
#             for run in paragraph.runs:
#                 run.font.name = 'DejaVu Sans'
#                 run.font.size = Pt(40)
#                 run.bold = True
#                 run.italic = True
#                 run.underline = True
#                 run.font.color.rgb = RGBColor(0, 0, 128)
#             # Optional: Center-align the paragraph
#             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#         elif text_to_find2 in paragraph.text:
#             # Replace the text with the new value
#             paragraph.text = paragraph.text.replace(text_to_find2, text2)
#
#             # Format the text
#             for run in paragraph.runs:
#                 run.font.name = 'DejaVu Sans'
#                 run.font.size = Pt(24)
#                 run.bold = True
#                 run.underline = False
#                 run.font.color.rgb = RGBColor(0, 0, 128)
#             # Optional: Center-align the paragraph
#             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#         elif text_to_find3 in paragraph.text:
#             # Replace the text with the new value
#             paragraph.text = paragraph.text.replace(text_to_find3, str(date))
#
#             # Format the text
#             for run in paragraph.runs:
#                 run.font.name = 'Cascadia Code SemiBold'
#                 run.font.size = Pt(16)
#                 run.bold = False
#                 run.underline = False
#                 run.font.color.rgb = RGBColor(0, 0, 255)
#
#             # Optional: Center-align the paragraph
#             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#
#     # Save the modified document
#     word_file_name = f"data/images/{user_id}{test_id}.docx"
#     pdf_file_name = f"data/images/{user_id}{test_id}.pdf"
#     doc.save(word_file_name)
#     convert(word_file_name, pdf_file_name)
#
#     try:
#         os.remove(word_file_name)
#         # print(f"{word_file_name} o'chirildi")
#     except OSError as xatolik:
#         print(f"{word_file_name} o'chirishda xatolik yuz berdi: {xatolik}")
#     return pdf_file_name


import os
import subprocess
import shutil  # Import the shutil module for file operations

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor


def create_certificate(user_id, test_id, fullname, science, class_num, date):
    # Load the Word document
    doc = Document("data/images/certificate.docx")

    # Define the text to replace and the replacement text
    text_to_find1 = "fullname"
    text_to_find2 = "science fani class_num-sinf testini"
    text_to_find3 = "date_time"

    text2 = text_to_find2.replace('science', science).replace('class_num', f"{class_num}")

    # Iterate through paragraphs in the document
    for paragraph in doc.paragraphs:
        if text_to_find1 in paragraph.text:
            # Replace the text with the new value
            paragraph.text = paragraph.text.replace(text_to_find1, fullname)

            # Format the text
            for run in paragraph.runs:
                run.font.name = 'DejaVu Sans'
                run.font.size = Pt(35)
                run.bold = True
                run.italic = True
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 128)
            # Optional: Center-align the paragraph
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        elif text_to_find2 in paragraph.text:
            # Replace the text with the new value
            paragraph.text = paragraph.text.replace(text_to_find2, text2)

            # Format the text
            for run in paragraph.runs:
                run.font.name = 'DejaVu Sans'
                run.font.size = Pt(20)
                run.bold = True
                run.underline = False
                run.font.color.rgb = RGBColor(0, 0, 128)
            # Optional: Center-align the paragraph
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        elif text_to_find3 in paragraph.text:
            # Replace the text with the new value
            paragraph.text = paragraph.text.replace(text_to_find3, str(date))

            # Format the text
            for run in paragraph.runs:
                run.font.name = 'Cascadia Code SemiBold'
                run.font.size = Pt(16)
                run.bold = False
                run.underline = False
                run.font.color.rgb = RGBColor(0, 0, 255)

            # Optional: Center-align the paragraph
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Output directory for the PDF file
    output_directory = 'data/images'

    # Save the modified document
    word_file_name = f"data/images/{user_id}{test_id}.docx"
    pdf_file_name = f"data/images/{user_id}{test_id}.pdf"
    doc.save(word_file_name)

    # Use unoconv to convert the DOCX to PDF
    try:
        subprocess.run(["unoconv", "-f", "pdf", word_file_name])
        os.remove(word_file_name)
    except Exception as xatolik:
        print(f"PDF ga o'girishda xatolik: {xatolik}")

    return pdf_file_name
